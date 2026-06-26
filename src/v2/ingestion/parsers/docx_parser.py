"""
docx_parser.py — Word Document Text Extraction for Dynamic Ingestion (v2)

Uses python-docx to extract text from .docx files.

Extraction strategy:
  - Iterate over all paragraphs in document order.
  - Iterate over all tables (cell text), placed after paragraphs.
  - Skip empty paragraphs.
  - Preserve heading levels with markdown-style prefixes (# ## ###)
    so the chunker can split on semantic boundaries.

Does NOT support:
  - .doc (legacy binary format) — advise users to save as .docx
  - Embedded images / charts — text-only extraction
  - Headers / footers — usually navigation boilerplate anyway
"""

from __future__ import annotations

import logging
from pathlib import Path
from typing import Union

import docx

from src.v2.ingestion.parsers.url_parser import ParsedDocument

logger = logging.getLogger(__name__)

# Map python-docx heading styles to markdown-style prefixes
_HEADING_PREFIXES: dict[str, str] = {
    "Heading 1": "# ",
    "Heading 2": "## ",
    "Heading 3": "### ",
    "Heading 4": "#### ",
    "Title": "# ",
    "Subtitle": "## ",
}


def _style_name(para) -> str:
    """Safely get the paragraph style name."""
    try:
        return para.style.name if para.style else ""
    except Exception:
        return ""


def parse_docx(
    source: Union[str, Path, bytes],
    filename: str = "document.docx",
) -> ParsedDocument:
    """
    Extract text from a Word .docx file.

    Args:
        source:   File path (str or Path) OR raw bytes from an UploadFile.
        filename: Human-readable filename used as document title fallback.

    Returns:
        ParsedDocument with structured text extracted from paragraphs and tables.

    Raises:
        ValueError: If the document contains no extractable text.
        RuntimeError: If python-docx cannot read the file.
    """
    try:
        if isinstance(source, bytes):
            import io
            doc = docx.Document(io.BytesIO(source))
        else:
            doc = docx.Document(str(source))
    except Exception as exc:
        raise RuntimeError(
            f"Failed to open Word document '{filename}': {exc}"
        ) from exc

    logger.info(f"Parsing Word document '{filename}'...")

    lines: list[str] = []
    title = filename

    # Extract paragraphs with heading-level prefixes
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style = _style_name(para)

        # Capture the first heading as the document title
        if style in ("Title", "Heading 1") and title == filename:
            title = text

        prefix = _HEADING_PREFIXES.get(style, "")
        lines.append(f"{prefix}{text}")

    # Extract table contents
    if doc.tables:
        lines.append("\n--- Tables ---")
        for table_idx, table in enumerate(doc.tables):
            lines.append(f"\n[Table {table_idx + 1}]")
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    lines.append(row_text)

    full_text = "\n".join(lines).strip()

    if not full_text:
        raise ValueError(
            f"No text could be extracted from '{filename}'. "
            "The document may be empty or contain only images."
        )

    logger.info(
        f"✓ Parsed Word document '{title}': {len(full_text)} chars, "
        f"{len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables"
    )

    return ParsedDocument(
        text=full_text,
        title=title,
        origin=str(source) if not isinstance(source, bytes) else filename,
        source_type="docx",
        page_count=None,  # python-docx doesn't expose page count easily
    )
